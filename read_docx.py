#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""读取财务指标示例.docx"""

from docx import Document

doc = Document('data_source/财务指标示例.docx')

print("=" * 60)
print("财务指标示例.docx 内容")
print("=" * 60)

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# 读取表格
for i, table in enumerate(doc.tables):
    print(f"\n【表格 {i+1}】")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
